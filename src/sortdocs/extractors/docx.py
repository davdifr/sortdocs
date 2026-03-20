from __future__ import annotations

from pathlib import Path

from docx import Document

from sortdocs.extractors.base import BaseExtractor
from sortdocs.models import ExtractedContent, ExtractedFileType


class DocxExtractor(BaseExtractor):
    file_type = ExtractedFileType.DOCX
    supported_extensions = (".docx",)

    def _extract(self, path: Path) -> ExtractedContent:
        document = Document(path)
        paragraphs: list[str] = []

        for paragraph in document.paragraphs:
            if not paragraph.text.strip():
                continue
            paragraphs.append(paragraph.text)
            if len("\n".join(paragraphs)) >= self.max_chars:
                break

        combined_text = "\n".join(paragraphs)
        properties = document.core_properties
        warnings: list[str] = []
        if not combined_text.strip():
            warnings.append("No readable body text was found in the DOCX file.")

        return ExtractedContent(
            file_type=self.file_type,
            title_guess=properties.title or self._guess_title(path=path, text=combined_text),
            plain_text_excerpt=combined_text,
            detected_language=None,
            metadata={
                "paragraph_count": len(document.paragraphs),
                "author": properties.author,
                "title": properties.title,
                "subject": properties.subject,
                "keywords": properties.keywords,
                "created": properties.created.isoformat() if properties.created else None,
            },
            extraction_warnings=warnings,
        )
