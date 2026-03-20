from __future__ import annotations

from pathlib import Path

from docx import Document
from PIL import Image
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from sortdocs.extractors import (
    DocxExtractor,
    FallbackExtractor,
    ImageExtractor,
    PdfExtractor,
    StubOCRBackend,
    TextExtractor,
    get_extractor,
)
from sortdocs.models import ExtractedFileType


def test_text_extractor_normalizes_whitespace_and_limits_excerpt(tmp_path: Path) -> None:
    path = tmp_path / "sample.md"
    path.write_text("#  Hello   World  \n\nThis   is   a   test.\n", encoding="utf-8")

    content = TextExtractor(max_chars=20).extract(path)

    assert content.file_type == ExtractedFileType.TEXT
    assert content.title_guess == "Hello World"
    assert content.plain_text_excerpt == "# Hello World\n\nThis…"
    assert content.metadata["encoding"] == "utf-8"


def test_pdf_extractor_extracts_text_and_metadata(tmp_path: Path) -> None:
    path = tmp_path / "sample.pdf"
    _create_pdf_with_text(path, "Hello PDF World")

    content = PdfExtractor(max_chars=200).extract(path)

    assert content.file_type == ExtractedFileType.PDF
    assert "Hello PDF World" in content.plain_text_excerpt
    assert content.metadata["page_count"] == 1
    assert content.extraction_warnings == []


def test_pdf_extractor_returns_warning_for_corrupt_file(tmp_path: Path) -> None:
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"not a pdf")

    content = PdfExtractor(max_chars=200).extract(path)

    assert content.file_type == ExtractedFileType.PDF
    assert content.plain_text_excerpt == ""
    assert content.extraction_warnings


def test_docx_extractor_extracts_main_text(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    document = Document()
    document.core_properties.title = "Quarterly Report"
    document.add_paragraph("Quarterly Report")
    document.add_paragraph("Revenue increased in Q1.")
    document.save(path)

    content = DocxExtractor(max_chars=200).extract(path)

    assert content.file_type == ExtractedFileType.DOCX
    assert content.title_guess == "Quarterly Report"
    assert "Revenue increased in Q1." in content.plain_text_excerpt
    assert content.metadata["paragraph_count"] == 2


def test_image_extractor_returns_stub_ocr_warning_and_metadata(tmp_path: Path) -> None:
    path = tmp_path / "sample.png"
    Image.new("RGB", (12, 8), color="white").save(path)

    content = ImageExtractor(max_chars=100, ocr_backend=StubOCRBackend()).extract(path)

    assert content.file_type == ExtractedFileType.IMAGE
    assert content.plain_text_excerpt == ""
    assert content.metadata["width"] == 12
    assert content.metadata["ocr"]["backend"] == "stub"
    assert content.extraction_warnings == [
        "OCR is not implemented yet for images; metadata-only extraction was used."
    ]


def test_fallback_extractor_never_crashes_and_get_extractor_uses_it(tmp_path: Path) -> None:
    path = tmp_path / "archive.bin"
    path.write_bytes(b"\x00\x01\x02")

    extractor = get_extractor(path, max_chars=100)
    content = extractor.extract(path)

    assert isinstance(extractor, FallbackExtractor)
    assert content.file_type == ExtractedFileType.FALLBACK
    assert content.title_guess == "archive"
    assert content.extraction_warnings


def _create_pdf_with_text(path: Path, text: str) -> None:
    writer = PdfWriter()
    page = writer.add_blank_page(width=300, height=300)
    stream = DecodedStreamObject()
    encoded_text = text.encode("latin-1", errors="replace")
    stream.set_data(b"BT /F1 18 Tf 40 150 Td (" + encoded_text + b") Tj ET")

    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font)
    resources = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {
                    NameObject("/F1"): font_ref,
                }
            )
        }
    )
    page[NameObject("/Resources")] = resources
    page[NameObject("/Contents")] = writer._add_object(stream)

    with path.open("wb") as handle:
        writer.write(handle)
